import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH


def clean_markdown(text: str) -> str:
    """
    清洗 Markdown 格式标记，使输出适合直接写入 Word 文档。
    - 去除 **加粗** 和 *斜体* 标记
    - 去除 # 标题标记
    - 去除无序列表 - 开头符号
    - 去除代码块标记 ```
    - 去除多余空行
    """
    # 去除代码块标记
    text = re.sub(r'```[\s\S]*?```', lambda m: m.group().replace('```', ''), text)
    text = text.replace('```', '')

    # 去除加粗和斜体标记
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)  # ***bold italic***
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)        # **bold**
    text = re.sub(r'\*(.*?)\*', r'\1', text)             # *italic*
    text = re.sub(r'__(.*?)__', r'\1', text)             # __bold__
    text = re.sub(r'_(.*?)_', r'\1', text)               # _italic_

    # 去除标题标记 (# ## ### etc.)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)

    # 去除无序列表标记 (- item / * item)
    text = re.sub(r'^[\-\*]\s+', '', text, flags=re.MULTILINE)

    # 去除有序列表多余标记（保留数字编号）
    # 例如 "1. " 保留为 "1. "

    # 去除链接标记 [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

    # 去除图片标记 ![alt](url) -> (图片: alt)
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'（图片：\1）', text)

    # 去除水平分割线
    text = re.sub(r'^[\-\*\_]{3,}\s*$', '', text, flags=re.MULTILINE)

    # 压缩多个连续空行为一个
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


class PatentDocGenerator:
    """生成符合 CNIPA 标准的专利文档"""

    def __init__(self):
        self.font_name = "仿宋_GB2312"
        self.font_name_fallback = "FangSong"
        self.font_size = Pt(12)   # 小四号
        self.line_spacing = Pt(28)  # 固定行距 28 磅

    def _apply_font(self, run):
        """为单个 Run 设置中文字体"""
        run.font.name = self.font_name
        run.font.size = self.font_size
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{self.font_name}"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn("w:eastAsia"), self.font_name)

    def _set_style(self, doc):
        """设置文档全局样式"""
        style = doc.styles["Normal"]
        font = style.font
        font.name = self.font_name
        font.size = self.font_size

        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{self.font_name}"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn("w:eastAsia"), self.font_name)

        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = self.line_spacing

    def _add_paragraph(self, doc, text, bold=False, alignment=None):
        """添加一个格式化段落"""
        p = doc.add_paragraph()
        if alignment is not None:
            p.alignment = alignment
        run = p.add_run(text)
        run.bold = bold
        self._apply_font(run)
        return p

    def generate_specification(self, title: str, content: str, output_path: str) -> str:
        """生成说明书 .docx"""
        doc = Document()
        self._set_style(doc)

        # 清洗 Markdown 标记
        content = clean_markdown(content)
        title = clean_markdown(title)

        # 标题
        self._add_paragraph(doc, title, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # 按段落拆分并添加 [000x] 编号
        counter = 1
        lines = content.split("\n")
        section_keywords = ["技术领域", "背景技术", "发明内容", "有益效果", "附图说明", "具体实施方式"]

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 跳过与标题重复的行
            if stripped == title.strip():
                continue

            # 检测章节标题（不编号）
            is_section = any(kw in stripped for kw in section_keywords)
            if is_section:
                self._add_paragraph(doc, stripped, bold=True)
            else:
                num_str = f"[{counter:04d}]"
                self._add_paragraph(doc, f"{num_str} {stripped}")
                counter += 1

        doc.save(output_path)
        return output_path

    def generate_claims(self, claims_text: str, output_path: str) -> str:
        """生成权利要求书 .docx"""
        doc = Document()
        self._set_style(doc)

        # 清洗 Markdown 标记
        claims_text = clean_markdown(claims_text)

        self._add_paragraph(doc, "权利要求书", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        for line in claims_text.split("\n"):
            stripped = line.strip()
            if stripped:
                self._add_paragraph(doc, stripped)

        doc.save(output_path)
        return output_path

    def generate_abstract(self, abstract_text: str, output_path: str) -> str:
        """生成说明书摘要 .docx"""
        doc = Document()
        self._set_style(doc)

        # 清洗 Markdown 标记
        abstract_text = clean_markdown(abstract_text)

        self._add_paragraph(doc, "说明书摘要", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._add_paragraph(doc, abstract_text.strip())

        doc.save(output_path)
        return output_path


generator = PatentDocGenerator()
